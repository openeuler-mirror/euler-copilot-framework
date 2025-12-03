import logging
from typing import Optional
from docx import Document as DocxDocument

logger = logging.getLogger(__name__)


def parse_docx(file_path: str) -> Optional[str]:
    """
    解析 DOCX 文件
    :param file_path: 文件路径
    :return: 文件内容
    """
    try:
        doc = DocxDocument(file_path)
        if not doc:
            logger.error("[DocParser] 无法打开docx文件")
            return None
        
        paragraphs = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                paragraphs.append(paragraph.text)
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text)
        
        content = '\n'.join(paragraphs)
        return content
    except Exception as e:
        logger.exception(f"[DocParser] 解析DOCX文件失败: {e}")
        return None


def parse_doc(file_path: str) -> Optional[str]:
    """
    解析 DOC 文件（旧版 Word 格式）
    :param file_path: 文件路径
    :return: 文件内容
    """
    try:
        doc = DocxDocument(file_path)
        paragraphs = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                paragraphs.append(paragraph.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text)
        content = '\n'.join(paragraphs)
        return content
    except Exception:
        logger.warning("[DocParser] python-docx 不支持 DOC 格式，尝试其他方法")
        logger.warning("[DocParser] DOC 格式解析需要额外工具，当前仅支持 DOCX")
        return None

